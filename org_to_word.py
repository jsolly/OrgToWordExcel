import docx
from my_secrets import REGRESSION_DEVEXT_DBQA_GIS

def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
      c = docx.oxml.shared.OxmlElement('w:color')
      c.set(docx.oxml.shared.qn('w:val'), color)
      rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
      u = docx.oxml.shared.OxmlElement('w:u')
      u.set(docx.oxml.shared.qn('w:val'), 'none')
      rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def get_items_from_folder(gis_obj, folder, item_types=None) -> list: # folder=None returns the root folder
    items = gis_obj.users.me.items(folder=folder)

    if item_types:
        items = [item for item in items if item.type in item_types]

    return items

if __name__ == "__main__":

    word_document = docx.Document()
    section = word_document.sections[0] # A section is a page. If you set a setting for the first page, it applies to all the other pages.
    section.orientation = docx.enum.section.WD_ORIENT.LANDSCAPE # Make the pages landscape so we have more room
    section.top_margin, section.bottom_margin, section.left_margin, section.right_margin = 457200, 457200, 457200, 457200 # Reduce margins for more room
    
    new_width, new_height = section.page_height, section.page_width
    section.page_width, section.page_height = new_width, new_height

    FOLDERS = REGRESSION_DEVEXT_DBQA_GIS.users.me.folders # Get all of your AGOL folders

    for FOLDER in FOLDERS:
        word_document.add_heading(f"{FOLDER['title']}", level=1) # Set a heading with the title of the current folder

        ITEMS = get_items_from_folder(gis_obj=REGRESSION_DEVEXT_DBQA_GIS, folder=FOLDER) # Get all Items from the current AGOL fodler

        table = word_document.add_table(rows=1, cols=6, style="Light Grid")
        first_row_cells = table.rows[0].cells # Define the first row which will be the coloumn names
        first_row_cells[0].text = 'Item Name'
        first_row_cells[1].text = 'id'
        first_row_cells[2].text = 'Item Type'
        first_row_cells[3].text = "Item Details Page"
        first_row_cells[4].text = "Thumbnail"

        print(f"I have {len(ITEMS)} items to work on")

        for INDEX, ITEM in enumerate(ITEMS):
            print (f"I am working on Item {INDEX}")
            row_cells = table.add_row().cells
            row_cells[0].text = ITEM.title
            row_cells[1].text = f"id:{ITEM.id}"
            row_cells[2].text = ITEM.type
            item_details_link = f"{REGRESSION_DEVEXT_DBQA_GIS._url}/home/item.html?id={ITEM.id}"
            add_hyperlink(row_cells[3].paragraphs[0], item_details_link, 'Item Details Page', '0000EE', True)

    word_document.save(f"{REGRESSION_DEVEXT_DBQA_GIS.users.me.username}.docx") # Save the document as username.docx